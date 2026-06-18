from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path("/Users/hassaanahmed/WORK/ajl-tours-main")
LOGO = Path(
    "/var/folders/10/g4xhxjxx4s3fx812shm0395m0000gn/T/"
    "codex-clipboard-ac431ce7-6c8b-4a52-b5d9-61d0ab1bb9e2.png"
)
OUT = ROOT / "invoices" / "Vireonix_AJL_Tours_Invoice.docx"
PDF_OUT = ROOT / "invoices" / "Vireonix_AJL_Tours_Invoice.pdf"

BLUE = "0B3D91"
LIGHT_BLUE = "EAF2FF"
INK = "111827"
MUTED = "5B6472"
BORDER = "CBD5E1"
GREEN = "15803D"
COMPANY_PHONE = "+923000516286"
COMPANY_ADDRESS = "51 Nishtar Block, Allama Iqbal Town, Lahore"
FINAL_AMOUNT = "PKR 50,000"

FEATURES = [
    "Improved the tour browsing experience so customers can view Switzerland tours more clearly.",
    "Added and refined tour detail pages with better images, descriptions, pricing, and booking flow.",
    "Improved checkout, order summary, participant selection, and payment success/failure pages.",
    "Added group discount pricing so larger bookings can show correct discounted totals.",
    "Added currency support and exchange-rate handling for clearer price display.",
    "Added customer dashboard features including bookings, wishlist/favorites, and support-related records.",
    "Improved the admin dashboard for managing tours, orders, users, travel records, divisions, settings, and homepage banners.",
    "Added email/notification support for booking and customer communication workflows.",
    "Added multilingual website support structure for English, French, German, and Italian content.",
    "Improved SEO files, robots.txt, sitemap, and production frontend/backend configuration.",
]

FIXES = [
    "Fixed tour creation errors in the admin panel, including missing data and division validation problems.",
    "Fixed migration/routing issues that caused some admin tools to show not found errors.",
    "Fixed database connection handling with clearer errors and more reliable production behavior.",
    "Fixed tour ID consistency issues that affected links, checkout, and tour detail loading.",
    "Fixed CORS and deployment configuration issues between frontend and backend.",
    "Improved login/security flow for admin access and protected admin routes.",
    "Cleaned up duplicate backend logic and improved error handling so problems are easier to diagnose.",
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_fixed_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, size=10.5, color=INK, bold=False, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def paragraph(doc, text="", size=10.5, color=INK, bold=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold)
    return p


def heading(doc, text):
    p = paragraph(doc, text, size=13, color=BLUE, bold=True, before=8, after=5)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run_font(r, size=10.2, color=INK)
    return p


def kv_table(doc, rows, widths=(1900, 4100), header_fill=None):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for label, value in rows:
        cells = table.add_row().cells
        for idx, width in enumerate(widths):
            set_fixed_width(cells[idx], width)
            set_cell_border(cells[idx])
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if header_fill and idx == 0:
                set_cell_shading(cells[idx], header_fill)
        lp = cells[0].paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(label)
        set_run_font(lr, size=9.5, color=MUTED, bold=True)
        vp = cells[1].paragraphs[0]
        vp.paragraph_format.space_after = Pt(0)
        vr = vp.add_run(value)
        set_run_font(vr, size=9.8, color=INK)
    return table


def money_table(doc):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.autofit = False
    rows = [
        ("Professional service fee", FINAL_AMOUNT),
        ("Final amount payable", FINAL_AMOUNT),
    ]
    for index, (label, value) in enumerate(rows):
        cells = table.add_row().cells
        for cell, width in zip(cells, (3500, 2300)):
            set_fixed_width(cell, width)
            set_cell_border(cell, color=BORDER, size="6")
            set_cell_margins(cell, top=120, bottom=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if index == len(rows) - 1:
                set_cell_shading(cell, "DCFCE7")
        lp = cells[0].paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(label)
        set_run_font(lr, size=10.2 if index < len(rows) - 1 else 11.5, color=INK, bold=index == len(rows) - 1)
        rp = cells[1].paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rp.paragraph_format.space_after = Pt(0)
        rr = rp.add_run(value)
        set_run_font(rr, size=10.2 if index < len(rows) - 1 else 12, color=GREEN if index == len(rows) - 1 else INK, bold=True)


def make_invoice():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)

    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    left, right = header.rows[0].cells
    set_fixed_width(left, 5500)
    set_fixed_width(right, 3800)
    for cell in (left, right):
        set_cell_border(cell, color="FFFFFF", size="0")
        set_cell_margins(cell, top=0, bottom=0, start=0, end=0)

    lp = left.paragraphs[0]
    lp.paragraph_format.space_after = Pt(2)
    if LOGO.exists():
        lp.add_run().add_picture(str(LOGO), width=Inches(2.1))
    paragraph(left, "Software Development & Website Improvement Services", size=9.5, color=MUTED, after=1)
    paragraph(left, f"Phone: {COMPANY_PHONE}", size=8.8, color=MUTED, after=1)
    paragraph(left, f"Address: {COMPANY_ADDRESS}", size=8.8, color=MUTED, after=0)

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = rp.add_run("INVOICE")
    set_run_font(run, size=25, color=BLUE, bold=True)
    rp.paragraph_format.space_after = Pt(4)
    for text in ("Invoice No: VIR-AJL-2026-001", "Date: June 14, 2026", "Status: Payment Due"):
        p = right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        set_run_font(r, size=9.8, color=MUTED if "Status" not in text else GREEN, bold="Status" in text)

    paragraph(doc, "", after=2)

    info = doc.add_table(rows=1, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    c1, c2 = info.rows[0].cells
    for cell, width in zip((c1, c2), (4550, 4550)):
        set_fixed_width(cell, width)
        set_cell_border(cell, color=BORDER)
        set_cell_margins(cell, top=130, bottom=130, start=160, end=160)
        set_cell_shading(cell, LIGHT_BLUE)

    p = c1.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Billed To")
    set_run_font(r, size=10, color=BLUE, bold=True)
    paragraph(c1, "AJL Tours / Client", size=11, color=INK, bold=True, after=2)
    paragraph(c1, "Website and booking system improvement work", size=9.5, color=MUTED, after=0)

    p = c2.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("From")
    set_run_font(r, size=10, color=BLUE, bold=True)
    paragraph(c2, "Vireonix", size=11, color=INK, bold=True, after=2)
    paragraph(c2, f"Phone: {COMPANY_PHONE}", size=9.5, color=MUTED, after=1)
    paragraph(c2, f"Address: {COMPANY_ADDRESS}", size=9.5, color=MUTED, after=1)
    paragraph(c2, "Payment receiver: HASAAN AHMED", size=9.5, color=MUTED, after=0)

    heading(doc, "Work Completed")
    paragraph(
        doc,
        "The following items were completed to improve the AJL Tours website, booking flow, admin panel, "
        "customer experience, and production reliability.",
        size=10.2,
        color=INK,
        after=4,
    )

    for item in FEATURES:
        bullet(doc, item)

    heading(doc, "Bugs Fixed & Stability Improvements")
    for item in FIXES:
        bullet(doc, item)

    doc.add_page_break()
    heading(doc, "Invoice Summary")
    money_table(doc)

    heading(doc, "Payment Details")
    kv_table(
        doc,
        [
            ("Account Title", "HASAAN AHMED"),
            ("Account Number", "00300111298378"),
            ("IBAN", "PK44MEZN0000300111298378"),
            ("Bank", "MEEZAN DIGITAL CENTRE"),
        ],
        widths=(2200, 5200),
        header_fill="F8FAFC",
    )

    paragraph(
        doc,
        "Thank you. This invoice covers the completed improvement and bug-fix work for the AJL Tours website.",
        size=10.2,
        color=MUTED,
        before=10,
        after=0,
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(f"Vireonix | Invoice VIR-AJL-2026-001 | Amount Due: {FINAL_AMOUNT}")
    set_run_font(footer_run, size=8.5, color=MUTED)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


def make_pdf():
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "InvoiceBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=12.5,
        textColor=colors.HexColor(f"#{INK}"),
        spaceAfter=5,
    )
    small = ParagraphStyle(
        "Small",
        parent=body,
        fontSize=8.7,
        leading=11,
        textColor=colors.HexColor(f"#{MUTED}"),
    )
    h = ParagraphStyle(
        "Heading",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=12.5,
        leading=15,
        textColor=colors.HexColor(f"#{BLUE}"),
        spaceBefore=8,
        spaceAfter=6,
    )
    title = ParagraphStyle(
        "InvoiceTitle",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=25,
        leading=29,
        textColor=colors.HexColor(f"#{BLUE}"),
        alignment=TA_RIGHT,
        spaceAfter=4,
    )
    right_small = ParagraphStyle("RightSmall", parent=small, alignment=TA_RIGHT)
    center_footer = ParagraphStyle("CenterFooter", parent=small, alignment=TA_CENTER)

    def p(text, style=body):
        return Paragraph(text, style)

    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=letter,
        leftMargin=0.58 * inch,
        rightMargin=0.58 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
        title="Vireonix AJL Tours Invoice",
    )

    logo = Image(str(LOGO), width=1.95 * inch, height=1.95 * inch * 0.46) if LOGO.exists() else p("Vireonix", h)
    header_left = [
        logo,
        Spacer(1, 3),
        p("Software Development & Website Improvement Services", small),
        p(f"Phone: {COMPANY_PHONE}", small),
        p(f"Address: {COMPANY_ADDRESS}", small),
    ]
    header_right = [
        p("INVOICE", title),
        p("Invoice No: VIR-AJL-2026-001", right_small),
        p("Date: June 14, 2026", right_small),
        p("<font color='#15803D'><b>Status: Payment Due</b></font>", right_small),
    ]
    story = [
        Table(
            [[header_left, header_right]],
            colWidths=[3.75 * inch, 3.45 * inch],
            style=TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ]
            ),
        ),
        Spacer(1, 8),
    ]

    info = Table(
        [
            [
                [
                    p("<b><font color='#0B3D91'>Billed To</font></b>", body),
                    p("<b>AJL Tours / Client</b>", body),
                    p("Website and booking system improvement work", small),
                ],
                [
                    p("<b><font color='#0B3D91'>From</font></b>", body),
                    p("<b>Vireonix</b>", body),
                    p(f"Phone: {COMPANY_PHONE}", small),
                    p(f"Address: {COMPANY_ADDRESS}", small),
                    p("Payment receiver: HASAAN AHMED", small),
                ],
            ]
        ],
        colWidths=[3.5 * inch, 3.5 * inch],
        style=TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{LIGHT_BLUE}")),
                ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor(f"#{BORDER}")),
                ("INNERGRID", (0, 0), (-1, -1), 0.7, colors.HexColor(f"#{BORDER}")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 11),
                ("RIGHTPADDING", (0, 0), (-1, -1), 11),
                ("TOPPADDING", (0, 0), (-1, -1), 9),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
            ]
        ),
    )
    story.extend([info, Spacer(1, 10)])

    story.append(p("Work Completed", h))
    story.append(
        p(
            "The following items were completed to improve the AJL Tours website, booking flow, admin panel, "
            "customer experience, and production reliability.",
            body,
        )
    )
    for item in FEATURES:
        story.append(p(f"- {item}", body))

    story.append(p("Bugs Fixed & Stability Improvements", h))
    for item in FIXES:
        story.append(p(f"- {item}", body))

    story.append(PageBreak())
    story.append(p("Invoice Summary", h))
    summary = Table(
        [
            ["Professional service fee", FINAL_AMOUNT],
            [Paragraph("<b>Final amount payable</b>", body), Paragraph(f"<b><font color='#15803D'>{FINAL_AMOUNT}</font></b>", body)],
        ],
        colWidths=[3.95 * inch, 2.25 * inch],
        hAlign="RIGHT",
        style=TableStyle(
            [
                ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor(f"#{BORDER}")),
                ("INNERGRID", (0, 0), (-1, -1), 0.7, colors.HexColor(f"#{BORDER}")),
                ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor("#DCFCE7")),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9.8),
                ("ALIGN", (1, 0), (1, -1), "RIGHT"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 11),
                ("RIGHTPADDING", (0, 0), (-1, -1), 11),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        ),
    )
    story.extend([summary, Spacer(1, 6)])

    payment = Table(
        [
            ["Account Title", "HASAAN AHMED"],
            ["Account Number", "00300111298378"],
            ["IBAN", "PK44MEZN0000300111298378"],
            ["Bank", "MEEZAN DIGITAL CENTRE"],
        ],
        colWidths=[1.65 * inch, 4.45 * inch],
        style=TableStyle(
            [
                ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor(f"#{BORDER}")),
                ("INNERGRID", (0, 0), (-1, -1), 0.7, colors.HexColor(f"#{BORDER}")),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F8FAFC")),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9.5),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor(f"#{MUTED}")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        ),
    )
    story.append(KeepTogether([p("Payment Details", h), payment]))
    story.extend(
        [
            Spacer(1, 10),
            p(
                "Thank you. This invoice covers the completed improvement and bug-fix work for the AJL Tours website.",
                small,
            ),
            Spacer(1, 8),
            p(f"Vireonix | Invoice VIR-AJL-2026-001 | Amount Due: {FINAL_AMOUNT}", center_footer),
        ]
    )

    doc.build(story)


if __name__ == "__main__":
    make_invoice()
    make_pdf()
